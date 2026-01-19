import re
import argparse

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# Common bullet markers that appear in extracted PDF text
BULLET_CHARS = [
    "➣", "➤", "➢", "➔",
    "•", "◦", "·", "∙", "‣", "⁃",
    "▪", "■", "◼", "◾", "◻", "□",
    "-", "–", "—",
]

# Regex: optional leading spaces, then one of the bullet chars, then the bullet text
BULLET_RE = re.compile(rf"^\s*(?:{'|'.join(re.escape(c) for c in BULLET_CHARS)})\s+(.*\S)\s*$")


def set_aptos_12(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Aptos (Body)"
    style.font.size = Pt(12)

    rfonts = style.element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Aptos (Body)")
    rfonts.set(qn("w:hAnsi"), "Aptos (Body)")
    rfonts.set(qn("w:cs"), "Aptos (Body)")
    rfonts.set(qn("w:eastAsia"), "Aptos (Body)")


def is_footer_noise(line: str) -> bool:
    l = line.strip()
    if not l:
        return True
    if re.fullmatch(r"\d+", l):
        return True
    if "@" in l:
        return True
    if l.lower().startswith("further reading"):
        return True
    return False


def looks_like_heading(line: str) -> bool:
    l = line.strip()
    if not l:
        return False
    # Don't treat bullets as headings
    if BULLET_RE.match(l):
        return False

    if len(l) > 80:
        return False

    if l.endswith(":"):
        return True

    words = re.findall(r"[A-Za-z']+", l)
    if not words:
        return False

    if l.isupper() and len(l) <= 60:
        return True

    title_like = sum(1 for w in words if w[0].isupper()) / max(1, len(words))
    if title_like >= 0.7 and len(l) <= 70:
        return True

    return False


def add_bold_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Aptos (Body)"
    run.font.size = Pt(12)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """
    level 0 -> List Bullet
    level 1 -> List Bullet 2
    level 2 -> List Bullet 3
    """
    level = max(0, min(level, 2))
    style = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
    p = doc.add_paragraph(text, style=style)
    for r in p.runs:
        r.font.name = "Aptos (Body)"
        r.font.size = Pt(12)


def normalize_lines(text: str):
    lines = []
    for ln in text.splitlines():
        ln = re.sub(r"\s+", " ", ln.strip())
        if ln and not is_footer_noise(ln):
            lines.append(ln)
    return lines


def bullet_text(line: str):
    """
    Return the bullet content if line starts with a recognised bullet marker, else None.
    """
    m = BULLET_RE.match(line)
    if not m:
        return None
    return m.group(1).strip()


# ----------------------------
# COORDINATE-BASED BULLET LEVELS
# ----------------------------
def _is_bullet_start_text(t: str) -> bool:
    if not t:
        return False
    t = t.lstrip()
    return bool(t) and t[0] in BULLET_CHARS


def _extract_bullet_x_positions(page: fitz.Page) -> list[float]:
    """
    Returns a list of x-positions (bullet glyph x) for bullet lines on the page,
    in reading order (top-to-bottom, left-to-right).

    This uses page.get_text("dict") (coordinates) but does NOT change the text pipeline.
    """
    d = page.get_text("dict")
    hits = []

    # Walk blocks/lines/spans; sort lines by y then x for stable reading order.
    for block in d.get("blocks", []):
        if block.get("type") != 0:  # 0 = text block
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue

            # Build the line text in a conservative way
            line_text = "".join(s.get("text", "") for s in spans)
            if not _is_bullet_start_text(line_text):
                continue

            # Find the span where the first non-space char lives; use its x0
            bullet_x = None
            for s in spans:
                st = s.get("text", "")
                if not st:
                    continue
                # Skip spans that are only whitespace
                if not st.strip():
                    continue

                # The bullet glyph is the first non-space char of the full line
                # If this span begins (after lstrip) with a bullet, take its x0.
                if st.lstrip() and st.lstrip()[0] in BULLET_CHARS:
                    bullet_x = float(s["bbox"][0])
                    break

                # Otherwise, sometimes the bullet is glued after spaces in the same span
                # We still treat the first non-space char as the bullet. If it is bullet, use this span x0.
                first = st.lstrip()[0] if st.lstrip() else ""
                if first in BULLET_CHARS:
                    bullet_x = float(s["bbox"][0])
                    break

            if bullet_x is None:
                # Fallback: use line bbox x0 if span parsing fails
                bbox = line.get("bbox")
                if bbox:
                    bullet_x = float(bbox[0])
                else:
                    continue

            bbox = line.get("bbox") or spans[0].get("bbox")
            y0 = float(bbox[1]) if bbox else 0.0
            x0 = float(bbox[0]) if bbox else bullet_x
            hits.append((y0, x0, bullet_x))

    hits.sort(key=lambda t: (t[0], t[1]))
    return [bx for _, __, bx in hits]


def _cluster_x_positions(xs: list[float], tol: float = 4.0) -> list[float]:
    """
    Cluster x positions into columns using a simple tolerance (points).
    Returns cluster centers sorted ascending.
    Deterministic and robust for slide bullets.
    """
    if not xs:
        return []
    xs_sorted = sorted(xs)
    clusters = [[xs_sorted[0]]]
    for x in xs_sorted[1:]:
        if abs(x - clusters[-1][-1]) <= tol:
            clusters[-1].append(x)
        else:
            clusters.append([x])

    centers = [sum(c) / len(c) for c in clusters]
    centers.sort()
    return centers


def _levels_for_bullets_on_page(bullet_xs: list[float]) -> list[int]:
    """
    Convert bullet x positions to levels 0/1/2 by clustering.
    Leftmost cluster => level 0, next => level 1, etc.
    """
    if not bullet_xs:
        return []
    centers = _cluster_x_positions(bullet_xs, tol=4.0)
    # Map each bullet x to the nearest cluster center index
    levels = []
    for x in bullet_xs:
        idx = min(range(len(centers)), key=lambda i: abs(x - centers[i]))
        levels.append(max(0, min(idx, 2)))
    return levels


# ----------------------------
# POSTPROCESSING ONLY: make bullets look like Word default (filled dot at all levels)
# and make headings bulleted + indent their children appropriately.
# (NO content changes.)
# ----------------------------
def postprocess_formatting(docx_path: str) -> None:
    """
    Postprocess the produced DOCX WITHOUT changing content:
    1) Remove bold-only heading paragraphs that have no bullet paragraphs beneath them (before next heading).
    2) Convert remaining headings into level-0 bullets ("List Bullet"), keeping the heading text bold.
    3) Force ALL bullets to use "List Bullet" (filled dot) and simulate nesting via indentation only.
       Bullets under a heading are shifted +1 level deeper (cap at 2).
    """

    doc = Document(docx_path)
    paras = doc.paragraphs

    def is_heading(p) -> bool:
        txt = (p.text or "").strip()
        if not txt:
            return False
        # don't treat list paragraphs as headings
        if p.style and p.style.name and p.style.name.startswith("List"):
            return False
        # Heading = all non-empty runs are bold
        any_nonempty_run = False
        for run in p.runs:
            if run.text and run.text.strip():
                any_nonempty_run = True
                if not run.bold:
                    return False
        return any_nonempty_run

    def is_bullet(p) -> bool:
        return bool(p.style and p.style.name and p.style.name.startswith("List"))

    def bullet_level_from_style_name(name: str) -> int:
        # Infer level from what the converter produced:
        # "List Bullet" => 0, "List Bullet 2" => 1, "List Bullet 3" => 2
        if name == "List Bullet":
            return 0
        if name == "List Bullet 2":
            return 1
        if name == "List Bullet 3":
            return 2
        return 0

    # Indentation parameters (points). Adjust if you want slightly different spacing,
    # but this keeps a clean Word-like nested bullet appearance.
    HANG = Pt(18)       # hanging indent so wrapped lines align after bullet glyph
    STEP = Pt(18)       # indent step per level
    BASE_LEFT = Pt(18)  # level 0 left indent

    def apply_level_indent(p, level: int) -> None:
        level = max(0, min(level, 2))
        pf = p.paragraph_format
        pf.left_indent = Pt(BASE_LEFT.pt + STEP.pt * level)
        pf.first_line_indent = Pt(-HANG.pt)

    def enforce_aptos_12(p) -> None:
        for r in p.runs:
            r.font.name = "Aptos (Body)"
            r.font.size = Pt(12)

    # ---- (1) Remove headings that have no bullets beneath them (same as your original cleanup) ----
    to_delete_idxs = []
    for i, p in enumerate(paras):
        if not is_heading(p):
            continue

        has_bullet = False
        for j in range(i + 1, len(paras)):
            nxt = paras[j]
            if is_heading(nxt):
                break
            if is_bullet(nxt):
                has_bullet = True
                break

        if not has_bullet:
            to_delete_idxs.append(i)

    for i in reversed(to_delete_idxs):
        p = paras[i]
        p._element.getparent().remove(p._element)

    # Refresh after deletions
    paras = doc.paragraphs

    # ---- (2)(3) Headings => List Bullet; bullets => List Bullet always, indent for nesting ----
    in_heading_block = False

    for p in paras:
        if is_heading(p):
            # Heading becomes a filled-dot bullet at level 0
            p.style = doc.styles["List Bullet"]
            apply_level_indent(p, 0)

            # Keep heading bold (already bold, but enforce)
            for r in p.runs:
                if r.text and r.text.strip():
                    r.bold = True
            enforce_aptos_12(p)

            in_heading_block = True
            continue

        if in_heading_block and is_bullet(p):
            # Shift bullets under heading +1 level, but keep filled-dot bullet glyph
            old_name = p.style.name if p.style else ""
            old_level = bullet_level_from_style_name(old_name)
            new_level = min(old_level + 1, 2)

            p.style = doc.styles["List Bullet"]
            apply_level_indent(p, new_level)
            enforce_aptos_12(p)
            continue

        # End heading block when we hit blank or non-bullet content
        if (p.text or "").strip() == "":
            in_heading_block = False
        else:
            if not is_bullet(p):
                in_heading_block = False

        # For bullets not under headings, still force filled-dot look and preserve their level
        if is_bullet(p):
            old_name = p.style.name if p.style else ""
            lvl = bullet_level_from_style_name(old_name)

            p.style = doc.styles["List Bullet"]
            apply_level_indent(p, lvl)
            enforce_aptos_12(p)
        else:
            enforce_aptos_12(p)

    doc.save(docx_path)


def convert(pdf_path: str, out_docx_path: str, force_all_lines_bullets: bool = False) -> None:
    """
    Original behaviour is unchanged when force_all_lines_bullets=False (default).

    New feature:
      - If force_all_lines_bullets=True, any non-heading line that is not explicitly a bullet
        is treated as its own bullet line at level 0.
    """
    pdf = fitz.open(pdf_path)
    doc = Document()
    set_aptos_12(doc)

    for page_index in range(pdf.page_count):
        page = pdf.load_page(page_index)

        # IMPORTANT: keep your existing text extraction so output content stays identical
        raw = page.get_text("text") or ""
        lines = normalize_lines(raw)
        if not lines:
            continue

        # Compute bullet levels using coordinates (does not affect 'lines')
        bullet_xs = _extract_bullet_x_positions(page)
        bullet_levels = _levels_for_bullets_on_page(bullet_xs)
        bullet_level_idx = 0

        # Slide title: first non-bullet line (or short fallback)
        title = None
        for ln in lines:
            if bullet_text(ln) is not None:
                continue
            if looks_like_heading(ln) or len(ln) <= 60:
                title = ln
                break

        # Skip slides titled Outline or Summary (case-insensitive)
        if title and title.strip().lower() in {"outline", "summary"}:
            continue

        if title:
            add_bold_line(doc, title)

        current_bullet = None
        current_level = 0

        def flush_bullet():
            nonlocal current_bullet, current_level
            if current_bullet:
                add_bullet(doc, current_bullet.strip(), current_level)
                current_bullet = None
                current_level = 0

        for ln in lines:
            if title and ln == title:
                continue

            bt = bullet_text(ln)
            is_head = looks_like_heading(ln)

            if bt is not None:
                # new bullet starts (original behaviour)
                flush_bullet()

                # Pull the next coordinate-derived level if available; else default to 0
                if bullet_level_idx < len(bullet_levels):
                    current_level = bullet_levels[bullet_level_idx]
                    bullet_level_idx += 1
                else:
                    current_level = 0

                current_bullet = bt
                continue

            if is_head:
                # heading line (original behaviour)
                flush_bullet()
                add_bold_line(doc, ln)
                continue

            # ---- NEW FEATURE (opt-in): treat every non-heading line as a bullet ----
            if force_all_lines_bullets:
                flush_bullet()
                current_level = 0
                current_bullet = ln
                continue

            # continuation line: append to existing bullet (original behaviour)
            if current_bullet:
                current_bullet += " " + ln

        flush_bullet()
        doc.add_paragraph("")

    # Save then postprocess formatting (ONLY formatting changes)
    doc.save(out_docx_path)
    postprocess_formatting(out_docx_path)
    print(f"Saved (formatted): {out_docx_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("pdf", help="Input slides PDF")
    ap.add_argument("out", help="Output docx")
    ap.add_argument(
        "--all-bullets",
        action="store_true",
        help="Treat every non-heading line as a bullet (useful for PDFs without bullet glyphs).",
    )
    args = ap.parse_args()
    convert(args.pdf, args.out, force_all_lines_bullets=args.all_bullets)


if __name__ == "__main__":
    main()
