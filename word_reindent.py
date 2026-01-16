import argparse
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


# ----------------------------
# STYLE / FONT
# ----------------------------
def set_aptos_12(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Aptos (Body)"
    style.font.size = Pt(12)

    rfonts = style.element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Aptos (Body)")
    rfonts.set(qn("w:hAnsi"), "Aptos (Body)")
    rfonts.set(qn("w:cs"), "Aptos (Body)")
    rfonts.set(qn("w:eastAsia"), "Aptos (Body)")


# ----------------------------
# DOCX XML HELPERS
# ----------------------------
def clear_document_content_keep_sectpr(doc: Document) -> None:
    """
    Remove all body children except the final sectPr (section properties),
    so the template's numbering/styles remain intact.
    """
    body = doc._element.body
    children = list(body)
    for child in children:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_paragraph_numbering(p, num_id: int, ilvl: int) -> None:
    """
    Force paragraph to use numbering definition (numId) and list level (ilvl).
    """
    ilvl = max(0, min(int(ilvl), 8))
    pPr = p._p.get_or_add_pPr()

    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement("w:numPr")

    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))

    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(int(num_id)))

    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def clear_direct_paragraph_formatting(p) -> None:
    """
    Avoid carrying manual indents/hanging indents/etc,
    because they can fight numbering.
    """
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.right_indent = None
    pf.space_before = None
    pf.space_after = None
    pf.line_spacing = None


def copy_runs_force_not_bold(src_p, dst_p) -> None:
    """
    Copy text and basic styling, but FORCE bold to normal (False).
    """
    if not src_p.runs:
        if src_p.text:
            r = dst_p.add_run(src_p.text)
            r.bold = False
            r.italic = None
            r.underline = None
            r.font.name = "Aptos (Body)"
            r.font.size = Pt(12)
        return

    for r0 in src_p.runs:
        if not r0.text:
            continue
        r = dst_p.add_run(r0.text)
        r.bold = False  # <-- force normal
        r.italic = r0.italic
        r.underline = r0.underline
        r.font.name = "Aptos (Body)"
        r.font.size = Pt(12)


# ----------------------------
# LIST DETECTION (FROM YOUR SOURCE DOC)
# ----------------------------
def is_list_like(src_p) -> bool:
    name = (src_p.style.name if src_p.style else "") or ""
    if name.startswith("List"):
        return True
    pf = src_p.paragraph_format
    return (pf.first_line_indent is not None and pf.first_line_indent.pt < 0) and (
        pf.left_indent is not None and pf.left_indent.pt > 0
    )


def infer_level_from_indent(src_p) -> int:
    """
    Your generated docs typically use:
      level 0 left indent ≈ 18pt
      level 1 left indent ≈ 36pt
      level 2 left indent ≈ 54pt
    """
    pf = src_p.paragraph_format
    if pf.left_indent is None:
        return 0
    left = float(pf.left_indent.pt)
    lvl = round((left - 18.0) / 18.0)
    return max(0, min(int(lvl), 2))


# ----------------------------
# MAIN CONVERSION
# ----------------------------
def apply_template_bullets(source_docx: str, template_docx: str, out_docx: str) -> None:
    src = Document(source_docx)

    # Start output as a copy of template so we inherit its numbering/bullets exactly
    out = Document(template_docx)
    set_aptos_12(out)

    # Wipe template content but keep its styles/numbering definitions
    clear_document_content_keep_sectpr(out)

    # Keep the same behavior as your working version:
    TEMPLATE_NUM_ID = 1
    TEMPLATE_LIST_STYLE = "List Paragraph"

    for sp in src.paragraphs:
        txt = sp.text or ""

        # ---- Change requested: remove weird empty spaces/newlines ----
        # Skip empty/whitespace-only paragraphs entirely.
        if not txt.strip():
            continue

        if is_list_like(sp):
            lvl = infer_level_from_indent(sp)

            dp = out.add_paragraph("")
            if TEMPLATE_LIST_STYLE in [s.name for s in out.styles]:
                dp.style = out.styles[TEMPLATE_LIST_STYLE]

            set_paragraph_numbering(dp, TEMPLATE_NUM_ID, lvl)
            clear_direct_paragraph_formatting(dp)

            # ---- Change requested: all bold -> normal ----
            copy_runs_force_not_bold(sp, dp)

        else:
            dp = out.add_paragraph("")
            clear_direct_paragraph_formatting(dp)

            # ---- Change requested: all bold -> normal ----
            copy_runs_force_not_bold(sp, dp)

        # Ensure font consistency + bold off (belt & braces)
        for r in dp.runs:
            r.font.name = "Aptos (Body)"
            r.font.size = Pt(12)
            r.bold = False

    out.save(out_docx)
    print(f"Saved: {out_docx}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("source", help="Your current generated notes DOCX")
    ap.add_argument("template", help="Template DOCX whose bullet formatting you want")
    ap.add_argument("out", help="Output DOCX")
    args = ap.parse_args()

    apply_template_bullets(args.source, args.template, args.out)


if __name__ == "__main__":
    main()
